"""
Resume Tailor Routes
API endpoints for AI-powered resume tailoring and optimization.
"""
import logging
from flask import Blueprint, request, jsonify, g, Response
from pydantic import ValidationError

from app import db
from app.models.candidate import Candidate
from app.models.job_posting import JobPosting
from app.models.candidate_job_match import CandidateJobMatch
from app.models.tailored_resume import TailoredResume, TailoredResumeStatus
from app.services.resume_tailor import ResumeTailorOrchestrator
from app.schemas.tailored_resume_schema import (
    TailorResumeRequest,
    TailorResumeFromMatchRequest,
    TailorManualResumeRequest,
    TailoredResumeResponse,
    TailoredResumeDetailResponse,
    TailoredResumeListResponse,
    ExportResumeRequest,
    ExportFormat
)
from app.middleware.portal_auth import require_portal_auth, require_permission
from app.middleware.tenant_context import with_tenant_context

logger = logging.getLogger(__name__)

resume_tailor_bp = Blueprint('resume_tailor', __name__, url_prefix='/api/resume-tailor')


def error_response(message: str, status: int = 400, details: dict = None):
    """Helper to create error responses"""
    response = {
        'error': 'Error',
        'message': message,
        'status': status
    }
    if details:
        response['details'] = details
    return jsonify(response), status


# ============================================================================
# Tailoring Endpoints
# ============================================================================

@resume_tailor_bp.route('/tailor', methods=['POST'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def tailor_resume():
    """
    Start resume tailoring for a candidate-job pair.
    
    POST /api/resume-tailor/tailor
    
    Request body:
    {
        "candidate_id": 123,
        "job_posting_id": 456,
        "target_score": 80,      // Optional, default 80
        "max_iterations": 3       // Optional, default 3
    }
    
    Returns:
        TailoredResumeResponse with initial status
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        
        # Validate request
        try:
            data = TailorResumeRequest.model_validate(request.get_json() or {})
        except ValidationError as e:
            return error_response("Validation error", 400, {"errors": e.errors()})
        
        # Verify candidate exists and belongs to tenant
        candidate = db.session.get(Candidate, data.candidate_id)
        if not candidate:
            return error_response(f"Candidate {data.candidate_id} not found", 404)
        if candidate.tenant_id != tenant_id:
            return error_response("Access denied", 403)
        
        # Verify job posting exists
        job_posting = db.session.get(JobPosting, data.job_posting_id)
        if not job_posting:
            return error_response(f"Job posting {data.job_posting_id} not found", 404)
        
        # Start tailoring
        orchestrator = ResumeTailorOrchestrator()
        
        tailored_resume = orchestrator.tailor_resume(
            candidate_id=data.candidate_id,
            job_posting_id=data.job_posting_id,
            tenant_id=tenant_id,
            target_score=data.target_score,
            max_iterations=data.max_iterations,
            resume_id=data.resume_id
        )
        
        return jsonify(tailored_resume.to_dict()), 201
        
    except ValueError as e:
        return error_response(str(e), 400)
    except Exception as e:
        logger.error(f"Resume tailoring failed: {e}")
        return error_response("Failed to tailor resume", 500, {"error": str(e)})


@resume_tailor_bp.route('/tailor-from-match', methods=['POST'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def tailor_from_match():
    """
    Start resume tailoring from an existing candidate-job match.
    
    POST /api/resume-tailor/tailor-from-match
    
    Request body:
    {
        "match_id": 123,
        "target_score": 80,      // Optional
        "max_iterations": 3       // Optional
    }
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        
        # Validate request
        try:
            data = TailorResumeFromMatchRequest.model_validate(request.get_json() or {})
        except ValidationError as e:
            return error_response("Validation error", 400, {"errors": e.errors()})
        
        # Verify match exists
        match = db.session.get(CandidateJobMatch, data.match_id)
        if not match:
            return error_response(f"Match {data.match_id} not found", 404)
        
        # Verify tenant access via candidate
        candidate = db.session.get(Candidate, match.candidate_id)
        if not candidate or candidate.tenant_id != tenant_id:
            return error_response("Access denied", 403)
        
        # Start tailoring
        orchestrator = ResumeTailorOrchestrator()
        
        tailored_resume = orchestrator.tailor_from_match(
            match_id=data.match_id,
            tenant_id=tenant_id,
            target_score=data.target_score,
            max_iterations=data.max_iterations
        )
        
        return jsonify(tailored_resume.to_dict()), 201
        
    except ValueError as e:
        return error_response(str(e), 400)
    except Exception as e:
        logger.error(f"Resume tailoring from match failed: {e}")
        return error_response("Failed to tailor resume", 500)


@resume_tailor_bp.route('/tailor-manual', methods=['POST'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def tailor_manual():
    """
    Start resume tailoring with a manually provided job description.
    
    This endpoint allows tailoring a resume without a job posting record.
    Recruiters can paste a job description directly.
    
    POST /api/resume-tailor/tailor-manual
    
    Request body:
    {
        "candidate_id": 123,
        "job_title": "Senior Software Engineer",
        "job_company": "Acme Corp",            // Optional
        "job_description": "We are looking for...",  // Required, min 50 chars
        "job_location": "New York, NY",        // Optional
        "target_score": 80,                    // Optional, default 80
        "max_iterations": 3                     // Optional, default 1
    }
    
    Returns:
        TailoredResumeResponse with results
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        
        # Validate request
        try:
            data = TailorManualResumeRequest.model_validate(request.get_json() or {})
        except ValidationError as e:
            return error_response("Validation error", 400, {"errors": e.errors()})
        
        # Verify candidate exists and belongs to tenant
        candidate = db.session.get(Candidate, data.candidate_id)
        if not candidate:
            return error_response(f"Candidate {data.candidate_id} not found", 404)
        if candidate.tenant_id != tenant_id:
            return error_response("Access denied", 403)
        
        # Start tailoring with manual job description
        orchestrator = ResumeTailorOrchestrator()
        
        tailored_resume = orchestrator.tailor_manual(
            candidate_id=data.candidate_id,
            job_title=data.job_title,
            job_description=data.job_description,
            tenant_id=tenant_id,
            job_company=data.job_company,
            job_location=data.job_location,
            target_score=data.target_score or 80,
            max_iterations=data.max_iterations or 1
        )
        
        return jsonify(tailored_resume.to_dict()), 201
        
    except ValueError as e:
        return error_response(str(e), 400)
    except Exception as e:
        logger.error(f"Manual resume tailoring failed: {e}")
        return error_response("Failed to tailor resume", 500, {"error": str(e)})


@resume_tailor_bp.route('/tailor-stream', methods=['POST'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def tailor_resume_stream():
    """
    Start resume tailoring with SSE progress streaming.
    
    POST /api/resume-tailor/tailor-stream
    
    Request body: Same as /tailor
    
    Returns:
        Server-Sent Events stream with progress updates
    
    Permissions: candidates.view
    """
    tenant_id = g.tenant_id
    
    # Validate request
    try:
        data = TailorResumeRequest.model_validate(request.get_json() or {})
    except ValidationError as e:
        return error_response("Validation error", 400, {"errors": e.errors()})
    
    # Verify access
    candidate = db.session.get(Candidate, data.candidate_id)
    if not candidate:
        return error_response(f"Candidate {data.candidate_id} not found", 404)
    if candidate.tenant_id != tenant_id:
        return error_response("Access denied", 403)
    
    job_posting = db.session.get(JobPosting, data.job_posting_id)
    if not job_posting:
        return error_response(f"Job posting {data.job_posting_id} not found", 404)
    
    def generate():
        """Generator for SSE events"""
        try:
            orchestrator = ResumeTailorOrchestrator()
            
            for event in orchestrator.tailor_resume_streaming(
                candidate_id=data.candidate_id,
                job_posting_id=data.job_posting_id,
                tenant_id=tenant_id,
                target_score=data.target_score,
                max_iterations=data.max_iterations,
                resume_id=data.resume_id
            ):
                yield event.to_sse()
                
        except Exception as e:
            logger.error(f"SSE streaming error: {e}")
            import json
            error_event = {
                'status': 'failed',
                'error': str(e),
                'message': 'Resume tailoring failed'
            }
            yield f"data: {json.dumps(error_event)}\n\n"
    
    return Response(
        generate(),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'Connection': 'keep-alive',
            'X-Accel-Buffering': 'no'  # Disable nginx buffering
        }
    )


# ============================================================================
# Retrieval Endpoints
# ============================================================================

@resume_tailor_bp.route('/<string:tailor_id>', methods=['GET'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def get_tailored_resume(tailor_id: str):
    """
    Get a tailored resume by ID.
    
    GET /api/resume-tailor/:tailor_id
    
    Query params:
        include_content: boolean - Include full resume markdown (default: false)
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        include_content = request.args.get('include_content', 'false').lower() == 'true'
        
        orchestrator = ResumeTailorOrchestrator()
        tailored_resume = orchestrator.get_tailored_resume(tailor_id)
        
        if not tailored_resume:
            return error_response(f"Tailored resume {tailor_id} not found", 404)
        
        # Verify tenant access
        if tailored_resume.tenant_id != tenant_id:
            return error_response("Access denied", 403)
        
        result = tailored_resume.to_dict()
        
        # Add related data
        if tailored_resume.candidate:
            result['candidate_name'] = (
                f"{tailored_resume.candidate.first_name} {tailored_resume.candidate.last_name}"
            )
        if tailored_resume.job_posting:
            result['job_title'] = tailored_resume.job_posting.title
            result['company'] = tailored_resume.job_posting.company
        
        # Remove content unless requested
        if not include_content:
            result.pop('original_resume_markdown', None)
            result.pop('tailored_resume_markdown', None)
        
        return jsonify(result), 200
        
    except Exception as e:
        logger.error(f"Failed to get tailored resume: {e}")
        return error_response("Failed to retrieve tailored resume", 500)


@resume_tailor_bp.route('/candidate/<int:candidate_id>', methods=['GET'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def get_candidate_tailored_resumes(candidate_id: int):
    """
    Get all tailored resumes for a candidate.
    
    GET /api/resume-tailor/candidate/:candidate_id
    
    Query params:
        status: Filter by status (pending, processing, completed, failed)
        page: Page number (default: 1)
        page_size: Items per page (default: 20)
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        
        # Verify candidate access
        candidate = db.session.get(Candidate, candidate_id)
        if not candidate:
            return error_response(f"Candidate {candidate_id} not found", 404)
        if candidate.tenant_id != tenant_id:
            return error_response("Access denied", 403)
        
        # Parse query params
        status_filter = request.args.get('status')
        page = int(request.args.get('page', 1))
        page_size = min(int(request.args.get('page_size', 20)), 100)
        
        # Build query
        from sqlalchemy import select
        
        stmt = (
            select(TailoredResume)
            .where(
                TailoredResume.candidate_id == candidate_id,
                TailoredResume.tenant_id == tenant_id
            )
        )
        
        if status_filter:
            try:
                status_enum = TailoredResumeStatus(status_filter)
                stmt = stmt.where(TailoredResume.status == status_enum)
            except ValueError:
                return error_response(f"Invalid status: {status_filter}", 400)
        
        stmt = stmt.order_by(TailoredResume.created_at.desc())
        
        # Paginate
        total = db.session.scalar(
            select(db.func.count()).select_from(stmt.subquery())
        )
        
        stmt = stmt.offset((page - 1) * page_size).limit(page_size)
        tailored_resumes = list(db.session.scalars(stmt))
        
        # Build response
        items = []
        for tr in tailored_resumes:
            item = tr.to_dict()
            # Remove content for list view
            item.pop('original_resume_markdown', None)
            item.pop('tailored_resume_markdown', None)
            
            if tr.job_posting:
                item['job_title'] = tr.job_posting.title
                item['company'] = tr.job_posting.company
            
            items.append(item)
        
        return jsonify({
            'items': items,
            'total': total,
            'page': page,
            'page_size': page_size,
            'total_pages': (total + page_size - 1) // page_size
        }), 200
        
    except Exception as e:
        logger.error(f"Failed to get candidate tailored resumes: {e}")
        return error_response("Failed to retrieve tailored resumes", 500)


# ============================================================================
# Comparison Endpoint
# ============================================================================

@resume_tailor_bp.route('/<string:tailor_id>/compare', methods=['GET'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def compare_resumes(tailor_id: str):
    """
    Get side-by-side comparison of original and tailored resume.
    
    GET /api/resume-tailor/:tailor_id/compare
    
    Returns:
        Original and tailored content with diff highlights
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        
        orchestrator = ResumeTailorOrchestrator()
        tailored_resume = orchestrator.get_tailored_resume(tailor_id)
        
        if not tailored_resume:
            return error_response(f"Tailored resume {tailor_id} not found", 404)
        
        if tailored_resume.tenant_id != tenant_id:
            return error_response("Access denied", 403)
        
        if tailored_resume.status != TailoredResumeStatus.COMPLETED:
            return error_response("Resume tailoring not completed yet", 400)
        
        return jsonify({
            'tailor_id': tailor_id,
            'original': {
                'content': tailored_resume.original_resume_content,
                'score': float(tailored_resume.original_match_score) * 100 if tailored_resume.original_match_score else None
            },
            'tailored': {
                'content': tailored_resume.tailored_resume_content,
                'score': float(tailored_resume.tailored_match_score) * 100 if tailored_resume.tailored_match_score else None
            },
            'improvements': tailored_resume.improvements,
            'matched_skills': tailored_resume.matched_skills or [],
            'missing_skills': tailored_resume.missing_skills or [],
            'added_skills': tailored_resume.added_skills or [],
            'score_improvement': float(tailored_resume.score_improvement) * 100 if tailored_resume.score_improvement else None
        }), 200
        
    except Exception as e:
        logger.error(f"Failed to compare resumes: {e}")
        return error_response("Failed to compare resumes", 500)


# ============================================================================
# Export Endpoint
# ============================================================================

@resume_tailor_bp.route('/<string:tailor_id>/export', methods=['GET'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def export_tailored_resume(tailor_id: str):
    """
    Export tailored resume in specified format.
    
    GET /api/resume-tailor/:tailor_id/export?format=pdf
    
    Query params:
        format: pdf | docx | markdown (default: pdf)
    
    Returns:
        File download
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        
        # Get tailored resume
        orchestrator = ResumeTailorOrchestrator()
        tailored_resume = orchestrator.get_tailored_resume(tailor_id)
        
        if not tailored_resume:
            return error_response(f"Tailored resume {tailor_id} not found", 404)
        
        if tailored_resume.tenant_id != tenant_id:
            return error_response("Access denied", 403)
        
        if tailored_resume.status != TailoredResumeStatus.COMPLETED:
            return error_response("Resume tailoring not completed yet", 400)
        
        if not tailored_resume.tailored_resume_content:
            return error_response("No tailored content available", 400)
        
        # Parse query params
        export_format = request.args.get('format', 'pdf')
        
        try:
            format_enum = ExportFormat(export_format)
        except ValueError:
            return error_response(f"Invalid format: {export_format}. Use: pdf, docx, markdown", 400)
        
        # Generate filename
        candidate_name = "resume"
        if tailored_resume.candidate:
            candidate_name = f"{tailored_resume.candidate.first_name}_{tailored_resume.candidate.last_name}"
        
        job_title = "tailored"
        if tailored_resume.job_title:
            job_title = tailored_resume.job_title.replace(" ", "_")[:30]
        
        filename = f"{candidate_name}_{job_title}_tailored"
        
        if format_enum == ExportFormat.MARKDOWN:
            # Return markdown directly
            return Response(
                tailored_resume.tailored_resume_content,
                mimetype='text/markdown',
                headers={
                    'Content-Disposition': f'attachment; filename="{filename}.md"'
                }
            )
        
        elif format_enum == ExportFormat.PDF:
            # Convert markdown to PDF using weasyprint
            try:
                from weasyprint import HTML, CSS
                import markdown
                import re
                
                logger.info(f"Starting PDF generation for tailor_id={tailor_id}")
                
                # Validate content exists
                content = tailored_resume.tailored_resume_content
                if not content:
                    logger.error(f"No tailored content for tailor_id={tailor_id}")
                    return error_response("No tailored content available for PDF export", 400)
                
                logger.info(f"Content length: {len(content)} chars")
                
                # ============================================================
                # CLEANUP: Fix common AI output issues
                # ============================================================
                
                # Remove duplicate sections (AI sometimes outputs same section twice)
                # Find all section headers and keep only the last occurrence of each
                section_pattern = r'^(#{1,3})\s+(.+?)$'
                lines = content.split('\n')
                seen_sections = {}
                section_ranges = []
                current_section_start = None
                current_section_header = None
                
                for i, line in enumerate(lines):
                    match = re.match(section_pattern, line.strip())
                    if match:
                        # Save previous section range
                        if current_section_header is not None:
                            section_ranges.append((current_section_start, i - 1, current_section_header))
                        current_section_start = i
                        current_section_header = match.group(2).strip().lower()
                
                # Don't forget the last section
                if current_section_header is not None:
                    section_ranges.append((current_section_start, len(lines) - 1, current_section_header))
                
                # Keep only the last occurrence of each section (which tends to be properly formatted)
                section_last_occurrence = {}
                for start, end, header in section_ranges:
                    # Normalize header for comparison (e.g., "Experience" and "## Experience" should match)
                    normalized_header = header.replace('#', '').strip().lower()
                    section_last_occurrence[normalized_header] = (start, end)
                
                # Build cleaned content by keeping only non-duplicate sections
                if section_ranges and len(section_ranges) > len(section_last_occurrence):
                    # There are duplicates - keep only last occurrences
                    lines_to_keep = set()
                    for header, (start, end) in section_last_occurrence.items():
                        for i in range(start, end + 1):
                            lines_to_keep.add(i)
                    
                    # Also keep any lines before the first section (name, contact info)
                    if section_ranges:
                        first_section_start = section_ranges[0][0]
                        for i in range(first_section_start):
                            lines_to_keep.add(i)
                    
                    # Rebuild content
                    cleaned_lines = [lines[i] for i in sorted(lines_to_keep)]
                    content = '\n'.join(cleaned_lines)
                    logger.info(f"Removed duplicate sections, new length: {len(content)} chars")
                
                # ============================================================
                # FIX FORMATTING: Ensure proper line breaks for markdown
                # ============================================================
                
                # 1. Ensure bullet points are on their own lines
                content = re.sub(r'([^\n])(\s*[-*]\s+)', r'\1\n\2', content)
                # 2. Ensure headers are on their own lines
                content = re.sub(r'([^\n])(#{1,3}\s+)', r'\1\n\n\2', content)
                # 3. Ensure there's a newline after headers before content
                content = re.sub(r'(#{1,3}\s+[^\n]+)(\n)([^#\n\-*])', r'\1\n\n\3', content)
                # 4. Fix inline bullet points that should be list items (e.g., "text - item - item")
                # Split lines that have multiple " - " patterns indicating inline bullets
                lines = content.split('\n')
                processed_lines = []
                for line in lines:
                    # Skip if already a proper bullet or header
                    if line.strip().startswith('-') or line.strip().startswith('*') or line.strip().startswith('#'):
                        processed_lines.append(line)
                    # Check for inline bullets pattern: " - " appearing multiple times
                    elif line.count(' - ') >= 2 and not line.strip().startswith('*'):
                        # This looks like inline bullet points, split them
                        # But first check if it's a header line like "Title | Company"
                        if ' | ' not in line:
                            parts = re.split(r'\s+-\s+', line)
                            if len(parts) > 1:
                                # First part might be a header/intro, rest are bullets
                                if parts[0].strip():
                                    processed_lines.append(parts[0].strip())
                                for part in parts[1:]:
                                    if part.strip():
                                        processed_lines.append(f"- {part.strip()}")
                            else:
                                processed_lines.append(line)
                        else:
                            processed_lines.append(line)
                    else:
                        processed_lines.append(line)
                
                content = '\n'.join(processed_lines)
                
                # Convert markdown to HTML
                html_content = markdown.markdown(
                    content,
                    extensions=['tables', 'fenced_code', 'nl2br']
                )
                
                logger.info(f"HTML content generated, length: {len(html_content)} chars")
                
                # Wrap with basic styling - 1 inch margins on all sides
                styled_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <style>
                        @page {{
                            size: letter;
                            margin: 1in;
                        }}
                        body {{
                            font-family: 'Helvetica Neue', Arial, sans-serif;
                            font-size: 11pt;
                            line-height: 1.4;
                            color: #333;
                            margin: 0;
                            padding: 0;
                        }}
                        h1 {{
                            font-size: 22pt;
                            margin-bottom: 5px;
                            margin-top: 0;
                            color: #1a1a1a;
                        }}
                        h1 + p {{
                            margin-top: 5px;
                            color: #555;
                        }}
                        h2 {{
                            font-size: 13pt;
                            border-bottom: 1.5px solid #2563eb;
                            padding-bottom: 4px;
                            margin-top: 18px;
                            margin-bottom: 8px;
                            color: #1a1a1a;
                            text-transform: uppercase;
                            letter-spacing: 0.5px;
                        }}
                        h3 {{
                            font-size: 11pt;
                            margin-bottom: 2px;
                            margin-top: 12px;
                            color: #333;
                            font-weight: 600;
                        }}
                        ul {{
                            margin: 5px 0;
                            padding-left: 18px;
                        }}
                        li {{
                            margin-bottom: 3px;
                        }}
                        p {{
                            margin: 4px 0;
                        }}
                        em {{
                            color: #555;
                            font-style: italic;
                        }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                
                # Generate PDF
                logger.info(f"Starting WeasyPrint PDF generation for tailor_id={tailor_id}")
                pdf_bytes = HTML(string=styled_html).write_pdf()
                logger.info(f"PDF generated successfully, size: {len(pdf_bytes)} bytes")
                
                return Response(
                    pdf_bytes,
                    mimetype='application/pdf',
                    headers={
                        'Content-Disposition': f'attachment; filename="{filename}.pdf"'
                    }
                )
                
            except ImportError as import_err:
                logger.error(f"WeasyPrint import failed: {import_err}")
                return error_response(
                    "PDF export not available. Install weasyprint: pip install weasyprint markdown",
                    501
                )
            except Exception as pdf_error:
                logger.error(f"PDF generation failed for tailor_id={tailor_id}: {pdf_error}", exc_info=True)
                return error_response(f"PDF generation failed: {str(pdf_error)}", 500)
        
        elif format_enum == ExportFormat.DOCX:
            # Convert to DOCX using python-docx
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from io import BytesIO
                import re
                
                doc = Document()
                
                # Set 1-inch margins on all sides
                for section in doc.sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)
                
                # Process markdown content
                lines = tailored_resume.tailored_resume_content.split('\n')
                is_first_line = True
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Headers
                    if line.startswith('# '):
                        p = doc.add_heading(line[2:], level=1)
                        # Style the main name header
                        if is_first_line:
                            for run in p.runs:
                                run.font.size = Pt(22)
                                run.font.bold = True
                        is_first_line = False
                    elif line.startswith('## '):
                        p = doc.add_heading(line[3:], level=2)
                        for run in p.runs:
                            run.font.size = Pt(13)
                    elif line.startswith('### '):
                        p = doc.add_heading(line[4:], level=3)
                        for run in p.runs:
                            run.font.size = Pt(11)
                    # Bullet points
                    elif line.startswith('- ') or line.startswith('* '):
                        p = doc.add_paragraph(line[2:], style='List Bullet')
                        for run in p.runs:
                            run.font.size = Pt(11)
                    # Regular paragraphs
                    else:
                        # Handle italic (dates)
                        if line.startswith('*') and line.endswith('*'):
                            p = doc.add_paragraph()
                            run = p.add_run(line.strip('*'))
                            run.italic = True
                            run.font.size = Pt(10)
                        # Contact info line (pipe separated)
                        elif ' | ' in line and is_first_line == False and not line.startswith('#'):
                            p = doc.add_paragraph(line)
                            for run in p.runs:
                                run.font.size = Pt(10)
                        else:
                            p = doc.add_paragraph(line)
                            for run in p.runs:
                                run.font.size = Pt(11)
                    
                    is_first_line = False
                
                # Save to bytes
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                return Response(
                    docx_buffer.getvalue(),
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    headers={
                        'Content-Disposition': f'attachment; filename="{filename}.docx"'
                    }
                )
                
            except ImportError:
                return error_response(
                    "DOCX export not available. Install python-docx: pip install python-docx",
                    501
                )
            except Exception as docx_error:
                logger.error(f"DOCX generation failed: {docx_error}")
                return error_response(f"DOCX generation failed: {str(docx_error)}", 500)
        
    except Exception as e:
        logger.error(f"Export failed: {e}")
        return error_response("Failed to export resume", 500)


# ============================================================================
# Stats Endpoint
# ============================================================================

@resume_tailor_bp.route('/stats', methods=['GET'])
@require_portal_auth
@with_tenant_context
@require_permission('candidates.view')
def get_tailoring_stats():
    """
    Get resume tailoring statistics for the tenant.
    
    GET /api/resume-tailor/stats
    
    Permissions: candidates.view
    """
    try:
        tenant_id = g.tenant_id
        from sqlalchemy import select, func
        
        # Total counts by status
        status_counts = db.session.execute(
            select(
                TailoredResume.status,
                func.count(TailoredResume.id)
            )
            .where(TailoredResume.tenant_id == tenant_id)
            .group_by(TailoredResume.status)
        ).all()
        
        counts = {s.value: 0 for s in TailoredResumeStatus}
        for status, count in status_counts:
            counts[status.value] = count
        
        # Average score improvement
        avg_improvement = db.session.scalar(
            select(
                func.avg(
                    TailoredResume.tailored_match_score - TailoredResume.original_match_score
                )
            )
            .where(
                TailoredResume.tenant_id == tenant_id,
                TailoredResume.status == TailoredResumeStatus.COMPLETED,
                TailoredResume.tailored_match_score.isnot(None),
                TailoredResume.original_match_score.isnot(None)
            )
        )
        
        # Average processing time
        avg_time = db.session.scalar(
            select(func.avg(TailoredResume.processing_time_seconds))
            .where(
                TailoredResume.tenant_id == tenant_id,
                TailoredResume.status == TailoredResumeStatus.COMPLETED,
                TailoredResume.processing_time_seconds.isnot(None)
            )
        )
        
        return jsonify({
            'total_tailored': sum(counts.values()),
            'by_status': counts,
            'successful': counts.get('completed', 0),
            'failed': counts.get('failed', 0),
            'average_score_improvement': round(float(avg_improvement), 1) if avg_improvement else None,
            'average_processing_time_seconds': round(float(avg_time), 1) if avg_time else None
        }), 200
        
    except Exception as e:
        logger.error(f"Failed to get stats: {e}")
        return error_response("Failed to get statistics", 500)
