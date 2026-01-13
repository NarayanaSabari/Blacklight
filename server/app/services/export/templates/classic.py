"""
Classic Resume Template

Traditional, formal design with serif fonts.
- Centered name and contact info
- Full-width underline section headers
- Skills displayed in categorized format
- Black and white only
"""

from io import BytesIO
from typing import List, Tuple
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .base_template import BaseResumeTemplate, TemplateMetadata


class ClassicTemplate(BaseResumeTemplate):
    """
    Classic resume template with traditional, formal design.
    
    Characteristics:
    - Serif font (Georgia, Times New Roman)
    - Centered header layout
    - Name: 26pt bold, centered
    - Section headers: 11pt uppercase, full-width underline
    - Skills: Categorized (Languages, Frameworks, Tools)
    - Black and white only
    - More traditional formatting with indented bullets
    """
    
    @property
    def metadata(self) -> TemplateMetadata:
        return TemplateMetadata(
            id="classic",
            name="Classic",
            description="Traditional, formal design with a professional feel. Best for corporate, finance, and legal roles.",
            is_default=False
        )
    
    def render_html(self, markdown_content: str) -> str:
        """Render markdown to styled HTML for PDF generation."""
        content = self._clean_markdown_content(markdown_content)
        html_body = self._markdown_to_html(content)
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{
            size: letter;
            margin: 1in 1in 1in 1in;
        }}
        
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: Georgia, 'Times New Roman', serif;
            font-size: 10.5pt;
            line-height: 1.5;
            color: #000;
        }}
        
        /* Name - H1 */
        h1 {{
            font-size: 26pt;
            font-weight: 700;
            text-align: center;
            margin-bottom: 4px;
            color: #000;
            text-transform: uppercase;
            letter-spacing: 2px;
        }}
        
        /* Contact info - paragraph after H1 */
        h1 + p {{
            font-size: 10pt;
            text-align: center;
            color: #000;
            margin-bottom: 6px;
        }}
        
        /* Second line of contact (if exists) */
        h1 + p + p {{
            font-size: 10pt;
            text-align: center;
            color: #000;
            margin-bottom: 16px;
        }}
        
        /* Section headers - H2 */
        h2 {{
            font-size: 11pt;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 1.5px;
            margin-top: 18px;
            margin-bottom: 10px;
            padding-bottom: 4px;
            border-bottom: 2px solid #000;
            color: #000;
        }}
        
        /* Job titles / subsections - H3 */
        h3 {{
            font-size: 11pt;
            font-weight: 700;
            margin-top: 12px;
            margin-bottom: 2px;
            color: #000;
            text-transform: uppercase;
        }}
        
        /* Company/dates line after H3 */
        h3 + p {{
            font-size: 10pt;
            color: #000;
            margin-bottom: 6px;
            font-style: italic;
        }}
        
        p {{
            margin: 4px 0;
            color: #000;
            text-align: justify;
        }}
        
        ul {{
            margin: 6px 0 10px 0;
            padding-left: 24px;
        }}
        
        li {{
            margin-bottom: 4px;
            color: #000;
            text-align: left;
        }}
        
        /* Italic text (dates, locations) */
        em {{
            font-style: italic;
        }}
        
        /* Bold text */
        strong {{
            font-weight: 700;
        }}
        
        /* Links */
        a {{
            color: #000;
            text-decoration: none;
        }}
        
        /* Horizontal rule for visual separation */
        hr {{
            border: none;
            border-top: 1px solid #000;
            margin: 12px 0;
        }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>"""
    
    def render_docx(self, markdown_content: str) -> bytes:
        """Render markdown to DOCX document."""
        content = self._clean_markdown_content(markdown_content)
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Process markdown lines
        lines = content.split('\n')
        is_first_heading = True
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # H1 - Name (centered, uppercase)
            if line.startswith('# '):
                text = line[2:].strip().upper()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(26)
                run.font.name = 'Georgia'
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.space_after = Pt(2)
                is_first_heading = False
                
            # H2 - Section headers (uppercase with border)
            elif line.startswith('## '):
                text = line[3:].strip().upper()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Georgia'
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.space_before = Pt(16)
                p.space_after = Pt(8)
                # Add thick bottom border
                self._add_bottom_border(p, thickness=12)
                
            # H3 - Job titles (uppercase)
            elif line.startswith('### '):
                text = line[4:].strip().upper()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Georgia'
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.space_before = Pt(10)
                p.space_after = Pt(1)
                
            # Bullet points (indented more)
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, text)
                p.paragraph_format.left_indent = Inches(0.35)
                p.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'Georgia'
                    
            # Regular paragraph
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = 'Georgia'
                    
                # Check if this looks like contact info (center it)
                if '|' in line or '@' in line or is_first_heading:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(10)
                    p.space_after = Pt(4)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _add_bottom_border(self, paragraph, thickness: int = 6):
        """Add a bottom border to a paragraph."""
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="{thickness}" w:space="1" w:color="000000"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with basic markdown formatting (bold, italic)."""
        # Pattern to match **bold** and *italic*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
