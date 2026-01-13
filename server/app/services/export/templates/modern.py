"""
Modern Resume Template

Clean, minimal design with sans-serif fonts.
- Left-aligned name and contact info
- Thin underline section headers
- Skills displayed as inline pipe-separated list
- Black and white only
"""

from io import BytesIO
from typing import List, Tuple
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .base_template import BaseResumeTemplate, TemplateMetadata


class ModernTemplate(BaseResumeTemplate):
    """
    Modern resume template with clean, minimal design.
    
    Characteristics:
    - Sans-serif font (Helvetica Neue, Arial)
    - Left-aligned layout
    - Name: 24pt bold
    - Section headers: 11pt uppercase with thin underline
    - Skills: Inline, pipe-separated
    - Black and white only
    """
    
    @property
    def metadata(self) -> TemplateMetadata:
        return TemplateMetadata(
            id="modern",
            name="Modern",
            description="Clean, minimal design with a contemporary feel. Best for tech, startups, and creative roles.",
            is_default=True
        )
    
    def render_html(self, markdown_content: str) -> str:
        """Render markdown to styled HTML for PDF generation."""
        content = self._clean_markdown_content(markdown_content)
        html_body = self._markdown_to_html(content)
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        @page {{
            size: letter;
            margin: 0.75in 0.75in 0.75in 0.75in;
        }}
        
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Helvetica Neue', Arial, sans-serif;
            font-size: 10pt;
            line-height: 1.4;
            color: #000;
        }}
        
        /* Name - H1 */
        h1 {{
            font-size: 24pt;
            font-weight: 700;
            margin-bottom: 4px;
            color: #000;
            letter-spacing: -0.5px;
        }}
        
        /* Contact info - paragraph after H1 */
        h1 + p {{
            font-size: 9pt;
            color: #333;
            margin-bottom: 16px;
            padding-bottom: 12px;
            border-bottom: 1px solid #000;
        }}
        
        /* Section headers - H2 */
        h2 {{
            font-size: 11pt;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 1px;
            margin-top: 16px;
            margin-bottom: 8px;
            padding-bottom: 4px;
            border-bottom: 1px solid #000;
            color: #000;
        }}
        
        /* Job titles / subsections - H3 */
        h3 {{
            font-size: 10pt;
            font-weight: 600;
            margin-top: 10px;
            margin-bottom: 2px;
            color: #000;
        }}
        
        /* Company/dates line after H3 */
        h3 + p {{
            font-size: 9pt;
            color: #333;
            margin-bottom: 4px;
        }}
        
        p {{
            margin: 4px 0;
            color: #000;
        }}
        
        ul {{
            margin: 4px 0 8px 0;
            padding-left: 16px;
        }}
        
        li {{
            margin-bottom: 2px;
            color: #000;
        }}
        
        /* Italic text (dates, locations) */
        em {{
            font-style: italic;
            color: #333;
        }}
        
        /* Bold text */
        strong {{
            font-weight: 600;
        }}
        
        /* Links */
        a {{
            color: #000;
            text-decoration: none;
        }}
        
        /* Tables for skills grid */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 8px 0;
        }}
        
        td {{
            padding: 2px 8px 2px 0;
            vertical-align: top;
        }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>"""
    
    def render_docx(self, markdown_content: str) -> bytes:
        """Render markdown to DOCX document."""
        content = self._clean_markdown_content(markdown_content)
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Process markdown lines
        lines = content.split('\n')
        is_first_heading = True
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # H1 - Name
            if line.startswith('# '):
                text = line[2:].strip()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(24)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.space_after = Pt(2)
                is_first_heading = False
                
            # H2 - Section headers
            elif line.startswith('## '):
                text = line[3:].strip().upper()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.space_before = Pt(14)
                p.space_after = Pt(6)
                # Add bottom border
                self._add_bottom_border(p)
                
            # H3 - Job titles
            elif line.startswith('### '):
                text = line[4:].strip()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.space_before = Pt(8)
                p.space_after = Pt(1)
                
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(p, text)
                p.paragraph_format.left_indent = Inches(0.25)
                p.space_after = Pt(1)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    
            # Regular paragraph
            else:
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                # Check if this looks like contact info (contains | or multiple items)
                if '|' in line or '@' in line:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(51, 51, 51)
                    p.space_after = Pt(8)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _add_bottom_border(self, paragraph):
        """Add a bottom border to a paragraph."""
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with basic markdown formatting (bold, italic)."""
        # Pattern to match **bold** and *italic*
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
