"""
Text extraction utilities for resume files (PDF, DOCX)
"""
import os
import tempfile
import traceback
from typing import Optional, Dict, Any, Tuple
import fitz  # PyMuPDF
from docx import Document
import pdfplumber
import logging
import unicodedata

from app.utils.docx_converter import DocxToPdfConverter, is_docx_conversion_enabled

logger = logging.getLogger(__name__)

# Magic bytes for file type validation
PDF_MAGIC = b'%PDF'
DOCX_MAGIC = b'PK'  # DOCX is ZIP-based


class TextExtractor:
    """
    Extract text from PDF and DOCX files
    Supports multiple extraction methods for better accuracy
    """
    
    @staticmethod
    def validate_file(file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate file integrity using magic bytes.
        
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}"
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return False, "File is empty (0 bytes)"
        
        if file_size < 10:
            return False, f"File too small ({file_size} bytes) - likely corrupted"
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
            
            if file_ext == '.pdf':
                if not header.startswith(PDF_MAGIC):
                    return False, "Invalid PDF file - does not start with %PDF header"
            elif file_ext in ['.docx', '.doc']:
                if not header.startswith(DOCX_MAGIC):
                    return False, "Invalid DOCX file - not a valid ZIP/DOCX format"
            
            return True, None
            
        except Exception as e:
            return False, f"Cannot read file: {str(e)}"
    
    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        Clean and normalize extracted text.
        Handles unicode issues and removes problematic characters.
        """
        if not text:
            return ""
        
        # Unicode normalization
        text = unicodedata.normalize('NFKD', text)
        
        # Remove null bytes and control characters (except newlines/tabs)
        text = ''.join(
            char for char in text 
            if char == '\n' or char == '\t' or (ord(char) >= 32 and ord(char) != 127)
        )
        
        # Encode/decode to handle any remaining encoding issues
        text = text.encode('utf-8', errors='ignore').decode('utf-8')
        
        return text
    
    @staticmethod
    def extract_from_file(file_path: str) -> Dict[str, Any]:
        """
        Extract text from a file (auto-detect type)
        
        Args:
            file_path: Path to resume file
        
        Returns:
            Dictionary with extracted text and metadata
        """
        logger.debug(f"[TextExtractor] Starting extraction from file: {file_path}")
        
        # Validate file first
        is_valid, error_msg = TextExtractor.validate_file(file_path)
        if not is_valid:
            logger.error(f"[TextExtractor] File validation failed: {error_msg}")
            raise ValueError(error_msg)
        
        file_ext = os.path.splitext(file_path)[1].lower()
        file_size = os.path.getsize(file_path)
        
        logger.debug(f"[TextExtractor] File details - extension: {file_ext}, size: {file_size} bytes ({file_size / 1024:.2f} KB)")
        
        result = None
        try:
            if file_ext == '.pdf':
                logger.debug(f"[TextExtractor] Detected PDF file, routing to extract_from_pdf")
                result = TextExtractor.extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                logger.debug(f"[TextExtractor] Detected DOCX/DOC file, routing to extract_from_docx_with_conversion")
                result = TextExtractor.extract_from_docx_with_conversion(file_path)
            else:
                logger.error(f"[TextExtractor] Unsupported file type: {file_ext}")
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Log final result summary
            text_length = len(result.get('text', '')) if result else 0
            method_used = result.get('method', 'unknown') if result else 'none'
            page_count = result.get('page_count', 0) if result else 0
            
            # Clean the extracted text
            if result and result.get('text'):
                result['text'] = TextExtractor.clean_extracted_text(result['text'])
                text_length = len(result['text'])
            
            logger.debug(f"[TextExtractor] Extraction complete - method: {method_used}, "
                        f"text_length: {text_length} chars, pages: {page_count}")
            
            return result
            
        except Exception as e:
            logger.error(f"[TextExtractor] Extraction failed for {file_path}: {str(e)}")
            logger.debug(f"[TextExtractor] Full traceback:\n{traceback.format_exc()}")
            raise
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF using multiple methods for best results
        
        Strategy:
        1. Try PyMuPDF (fast, good for simple PDFs)
        2. Fallback to pdfplumber (better for tables/columns)
        
        Returns:
            {
                'text': str,
                'page_count': int,
                'method': str,
                'has_images': bool,
                'metadata': dict
            }
        """
        logger.debug(f"[TextExtractor] Starting PDF extraction: {file_path}")
        
        result = {
            'text': '',
            'page_count': 0,
            'method': '',
            'has_images': False,
            'metadata': {}
        }
        
        # Method 1: PyMuPDF (primary)
        logger.debug(f"[TextExtractor] Attempting extraction with PyMuPDF")
        try:
            text = TextExtractor._extract_with_pymupdf(file_path)
            text_length = len(text) if text else 0
            stripped_length = len(text.strip()) if text else 0
            
            logger.debug(f"[TextExtractor] PyMuPDF extraction - raw_chars: {text_length}, "
                        f"stripped_chars: {stripped_length}")
            
            if text and len(text.strip()) > 50:  # Valid extraction
                logger.debug(f"[TextExtractor] PyMuPDF extraction successful (>{50} chars)")
                
                # Use context manager to prevent resource leaks
                with fitz.open(file_path) as doc:
                    result['text'] = text
                    result['page_count'] = len(doc)
                    result['method'] = 'pymupdf'
                    result['metadata'] = dict(doc.metadata) if doc.metadata else {}
                    
                    # Check for images
                    for page in doc:
                        if page.get_images():
                            result['has_images'] = True
                            break
                
                logger.debug(f"[TextExtractor] PyMuPDF result - pages: {result['page_count']}, "
                            f"has_images: {result['has_images']}, text_chars: {len(result['text'])}")
                return result
            else:
                logger.debug(f"[TextExtractor] PyMuPDF extraction insufficient "
                            f"({stripped_length} chars, need >50)")
        
        except Exception as e:
            logger.warning(f"[TextExtractor] PyMuPDF extraction failed: {str(e)}")
            logger.debug(f"[TextExtractor] PyMuPDF traceback:\n{traceback.format_exc()}")
        
        # Method 2: pdfplumber (fallback)
        logger.debug(f"[TextExtractor] Falling back to pdfplumber extraction")
        try:
            text = TextExtractor._extract_with_pdfplumber(file_path)
            text_length = len(text) if text else 0
            
            logger.debug(f"[TextExtractor] pdfplumber extraction - chars: {text_length}")
            
            if text:
                with pdfplumber.open(file_path) as pdf:
                    result['text'] = text
                    result['page_count'] = len(pdf.pages)
                    result['method'] = 'pdfplumber'
                    result['metadata'] = pdf.metadata
                
                logger.debug(f"[TextExtractor] pdfplumber extraction successful - "
                            f"pages: {result['page_count']}, text_chars: {len(result['text'])}")
                return result
            else:
                logger.warning(f"[TextExtractor] pdfplumber extraction returned empty text")
                raise RuntimeError(f"Failed to extract text from PDF: {file_path}")
        
        except Exception as e:
            logger.error(f"[TextExtractor] pdfplumber extraction failed: {str(e)}")
            logger.debug(f"[TextExtractor] pdfplumber traceback:\n{traceback.format_exc()}")
            raise RuntimeError(f"Failed to extract text from PDF: {file_path}")
    
    @staticmethod
    def _extract_with_pymupdf(file_path: str) -> str:
        """
        Extract text using PyMuPDF (fast method)
        Uses context manager to prevent resource leaks
        """
        text_parts = []
        
        with fitz.open(file_path) as doc:
            # Check for encrypted/password-protected PDF
            if doc.is_encrypted:
                raise ValueError("Password-protected PDF cannot be processed")
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def _extract_with_pdfplumber(file_path: str) -> str:
        """
        Extract text using pdfplumber (better for complex layouts)
        """
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract text
                text = page.extract_text()
                if text:
                    text_parts.append(text)
                
                # Extract tables separately
                tables = page.extract_tables()
                for table in tables:
                    # Convert table to text
                    table_text = '\n'.join([
                        ' | '.join([cell or '' for cell in row])
                        for row in table
                    ])
                    text_parts.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
        
        return '\n\n'.join(text_parts)
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file
        
        Returns:
            {
                'text': str,
                'page_count': int (estimated),
                'method': str,
                'has_images': bool,
                'metadata': dict
            }
        """
        result = {
            'text': '',
            'page_count': 0,
            'method': 'python-docx',
            'has_images': False,
            'metadata': {}
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    table_text.append(row_text)
                
                if table_text:
                    text_parts.append('\n[TABLE]\n' + '\n'.join(table_text) + '\n[/TABLE]\n')
            
            result['text'] = '\n\n'.join(text_parts)
            
            # Estimate page count (rough: 500 words per page)
            word_count = len(result['text'].split())
            result['page_count'] = max(1, word_count // 500)
            
            # Check for images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    result['has_images'] = True
                    break
            
            # Metadata
            core_props = doc.core_properties
            result['metadata'] = {
                'author': core_props.author,
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'title': core_props.title,
            }
            
            return result
        
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX: {file_path}") from e
    
    @staticmethod
    def extract_from_docx_with_conversion(file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX by converting to PDF first (better accuracy).
        Falls back to direct DOCX extraction if conversion fails.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            {
                'text': str,
                'page_count': int,
                'method': str,
                'has_images': bool,
                'metadata': dict
            }
        """
        logger.debug(f"[TextExtractor] Starting DOCX extraction with conversion: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        file_size = os.path.getsize(file_path)
        
        logger.debug(f"[TextExtractor] DOCX file details - extension: {file_ext}, size: {file_size} bytes ({file_size / 1024:.2f} KB)")
        
        if not is_docx_conversion_enabled():
            logger.info("[TextExtractor] DOCX conversion not available, using direct extraction")
            result = TextExtractor.extract_from_docx(file_path)
            result['method'] = f"{result['method']}_direct_fallback"
            logger.debug(f"[TextExtractor] Direct DOCX extraction complete - method: {result['method']}, "
                        f"text_chars: {len(result.get('text', ''))}")
            return result
        
        temp_pdf_path = None
        try:
            logger.info(f"[TextExtractor] Converting DOCX to PDF for better extraction: {file_path}")
            temp_pdf_path = DocxToPdfConverter.convert_to_temp_pdf(file_path)
            temp_pdf_size = os.path.getsize(temp_pdf_path)
            
            logger.debug(f"[TextExtractor] DOCX->PDF conversion successful - temp file: {temp_pdf_path}, "
                        f"size: {temp_pdf_size} bytes")
            
            logger.debug(f"[TextExtractor] Extracting text from converted PDF")
            result = TextExtractor.extract_from_pdf(temp_pdf_path)
            result['method'] = f"{result['method']}_from_converted_docx"
            result['source_format'] = 'docx'
            
            logger.info(f"[TextExtractor] Successfully extracted text via DOCX->PDF conversion "
                       f"(method: {result['method']}, text_chars: {len(result.get('text', ''))})")
            return result
        
        except Exception as e:
            logger.warning(f"[TextExtractor] DOCX to PDF conversion failed: {str(e)}. Falling back to direct DOCX extraction")
            logger.debug(f"[TextExtractor] Conversion failure traceback:\n{traceback.format_exc()}")
            
            logger.debug(f"[TextExtractor] Attempting direct DOCX extraction as fallback")
            result = TextExtractor.extract_from_docx(file_path)
            result['method'] = f"{result['method']}_direct_fallback"
            
            logger.debug(f"[TextExtractor] Direct DOCX fallback extraction complete - method: {result['method']}, "
                        f"text_chars: {len(result.get('text', ''))}")
            return result
        
        finally:
            if temp_pdf_path and os.path.exists(temp_pdf_path):
                try:
                    os.remove(temp_pdf_path)
                    logger.debug(f"[TextExtractor] Cleaned up temporary PDF: {temp_pdf_path}")
                except Exception as cleanup_error:
                    logger.warning(f"[TextExtractor] Failed to cleanup temp PDF {temp_pdf_path}: {cleanup_error}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text for better parsing
        
        - Remove excessive whitespace
        - Normalize line breaks
        - Remove special characters
        """
        # Remove multiple spaces
        text = ' '.join(text.split())
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n')
        text = text.replace('\r', '\n')
        
        # Remove excessive newlines (more than 2)
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text.strip()
    
    @staticmethod
    def extract_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract only metadata from file (no text extraction)
        Uses context managers to prevent resource leaks
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            try:
                with fitz.open(file_path) as doc:
                    metadata = {
                        'page_count': len(doc),
                        'metadata': dict(doc.metadata) if doc.metadata else {},
                        'file_size': os.path.getsize(file_path),
                        'is_encrypted': doc.is_encrypted,
                    }
                    return metadata
            except Exception as e:
                logger.warning(f"Failed to extract PDF metadata: {e}")
                return {'error': str(e)}
        
        elif file_ext in ['.docx', '.doc']:
            try:
                doc = Document(file_path)
                core_props = doc.core_properties
                
                return {
                    'metadata': {
                        'author': core_props.author,
                        'created': str(core_props.created) if core_props.created else None,
                        'modified': str(core_props.modified) if core_props.modified else None,
                        'title': core_props.title,
                    },
                    'file_size': os.path.getsize(file_path),
                }
            except Exception as e:
                logger.warning(f"Failed to extract DOCX metadata: {e}")
                return {'error': str(e)}
        
        return {'error': f'Unsupported file type: {file_ext}'}
